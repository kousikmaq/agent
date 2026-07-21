"""
Generates the FEATURES document (DOCX) with screenshots for the
Production Planning & Schedule Optimization Agent.

Run: python generate_features_doc.py
Output: Production_Planning_Agent_FEATURES.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1F, 0x2A, 0x37); NAVY = RGBColor(0x1B, 0x3A, 0x5E)
TEAL = RGBColor(0x0F, 0x76, 0x6E); SLATE = RGBColor(0x5A, 0x67, 0x78)
GREEN = RGBColor(0x2E, 0x7D, 0x46); AMBER = RGBColor(0xB4, 0x7A, 0x1E)
RED = RGBColor(0xC0, 0x39, 0x2B); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FILL_NAVY = "1B3A5E"; FILL_BLUE = "E8F0FA"; FILL_TEAL = "E2F3F1"
FILL_AMBER = "FCF3E3"; FILL_GREEN = "EAF6EC"; FILL_GREY = "F2F4F7"; FILL_MINT = "E6F4EC"
HDR = "1B3A5E"; BODY = "Segoe UI"; HEAD = "Segoe UI Semibold"; MONO = "Consolas"

ROOT = os.path.dirname(os.path.abspath(__file__))


def _shade(c, f):
    p = c._tc.get_or_add_tcPr(); s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), f); p.append(s)


def _mar(c, t=80, b=80, l=120, r=120):
    tcPr = c._tc.get_or_add_tcPr(); m = OxmlElement("w:tcMar")
    for side, v in (("top", t), ("bottom", b), ("start", l), ("end", r)):
        n = OxmlElement(f"w:{side}"); n.set(qn("w:w"), str(v)); n.set(qn("w:type"), "dxa"); m.append(n)
    tcPr.append(m)


def _border(t, color="D0D7DE"):
    b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}"); x.set(qn("w:val"), "single"); x.set(qn("w:sz"), "4")
        x.set(qn("w:space"), "0"); x.set(qn("w:color"), color); b.append(x)
    t._tbl.tblPr.append(b)


def _run(p, text, size=10.5, bold=False, italic=False, color=INK, font=BODY):
    r = p.add_run(text); r.font.name = font; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color; return r


def section(doc, num, title, fill=FILL_NAVY):
    t = doc.add_table(rows=1, cols=1); c = t.cell(0, 0); _shade(c, fill); _mar(c, 110, 110, 160, 160)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    _run(p, f"{num}   {title}", size=14, bold=True, color=WHITE, font=HEAD)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def subhead(doc, text, color=NAVY):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    _run(p, text, size=11.5, bold=True, color=color, font=HEAD)


def para(doc, text, size=10.5):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.12
    _run(p, text, size=size, color=INK); return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.1
        if isinstance(it, tuple):
            _run(p, it[0] + ": ", size=10.2, bold=True, color=NAVY); _run(p, it[1], size=10.2, color=INK)
        else:
            _run(p, it, size=10.2, color=INK)


def card(doc, title, body, fill=FILL_BLUE, accent=TEAL):
    t = doc.add_table(rows=1, cols=1); c = t.cell(0, 0); _shade(c, fill); _mar(c, 110, 110, 170, 170)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    _run(p, title, size=11, bold=True, color=accent, font=HEAD)
    b = c.add_paragraph(); b.paragraph_format.space_after = Pt(0); _run(b, body, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def table(doc, headers, rows, widths=None, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; _border(t)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; _shade(c, HDR); _mar(c)
        _run(c.paragraphs[0], h, size=size - 0.2, bold=True, color=WHITE, font=HEAD)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            _mar(cells[ci])
            if ri % 2 == 1:
                _shade(cells[ci], "F7F9FB")
            p = cells[ci].paragraphs[0]; p.paragraph_format.line_spacing = 1.05
            _run(p, str(val), size=size, bold=(ci == 0), color=(NAVY if ci == 0 else INK))
    if widths:
        for i, w in enumerate(widths):
            for cell in t.columns[i].cells:
                cell.width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def shot(doc, filename, width_in=6.3, caption=None):
    path = os.path.join(ROOT, filename)
    if not os.path.exists(path):
        para(doc, f"[screenshot missing: {filename}]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    # frame with a thin border via a 1-cell table
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; _border(tbl, "C7D2DE")
    cell = tbl.cell(0, 0); _mar(cell, 40, 40, 40, 40)
    cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run().add_picture(path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6); _run(cap, caption, size=8.8, italic=True, color=SLATE)


def mono(doc, lines, fill="1E2A38"):
    t = doc.add_table(rows=1, cols=1); c = t.cell(0, 0); _shade(c, fill); _mar(c, 100, 100, 140, 140)
    for i, ln in enumerate(lines):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
        _run(p, ln, size=8.8, color=RGBColor(0xE6, 0xED, 0xF3), font=MONO)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def feature(doc, num, title, fill, accent, what, how, criterion, image, caption):
    section(doc, num, title, fill=fill)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    _run(p, "What it does.  ", size=10.5, bold=True, color=accent); _run(p, what, size=10.5, color=INK)
    ph = doc.add_paragraph(); ph.paragraph_format.space_after = Pt(1)
    _run(ph, "How it works", size=10.5, bold=True, color=accent)
    bullets(doc, how)
    pc = doc.add_paragraph(); pc.paragraph_format.space_after = Pt(4)
    _run(pc, "Acceptance criterion.  ", size=10.2, bold=True, color=GREEN); _run(pc, criterion, size=10.2, italic=True, color=INK)
    shot(doc, image, caption=caption)
    doc.add_page_break()


def build():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = BODY; n.font.size = Pt(10.5); n.font.color.rgb = INK
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], m, Inches(0.7))

    # cover
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "FEATURE GUIDE  ·  WITH SCREENSHOTS", size=12, bold=True, color=TEAL)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Production Planning & Schedule\nOptimization Agent", size=28, bold=True, color=NAVY, font=HEAD)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "A working app: FastAPI + OR-Tools + React/Recharts, with an Azure OpenAI copilot.", size=13, color=SLATE)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "9 features · 8 dashboard sections · 500+ order connected dataset · 10 automated tests", size=11, italic=True, color=SLATE)
    doc.add_page_break()

    # 1 overview
    section(doc, "1", "Overview & how to read this guide")
    para(doc,
         "This agent turns a live order book into a feasible, optimized, explainable production plan. "
         "It is built on a strict principle: a deterministic engine (Operations Research) computes every "
         "number; the AI copilot only interprets intent, explains results and recommends actions - it never "
         "invents figures. Each feature below includes what it does, how it works, the acceptance criterion "
         "it satisfies, and a screenshot of the live app.")
    card(doc, "The running example",
         "A valve factory with 9 machines (work centres), 45 products, 520 sales orders across a 16-week "
         "horizon. Three weeks are overloaded, which the features detect, explain and resolve.",
         fill=FILL_TEAL, accent=TEAL)
    subhead(doc, "Acceptance criteria coverage (all met)")
    table(doc, ["Criterion", "Feature", "Status"],
          [["Capacity bottlenecks identified correctly", "Capacity Analysis", "Done"],
           ["Scheduling conflicts highlighted", "Capacity overloads + CP-SAT no-overlap", "Done"],
           ["Optimized production plan generated", "Schedule Optimization (OR-Tools)", "Done"],
           ["At-risk orders identified", "Delay Risk", "Done"],
           ["Alternative scheduling recommendations", "Resource Allocation + Scenarios", "Done"],
           ["Supports minimum 3 planning scenarios", "What-if Scenarios (3)", "Done"]],
          widths=[3.1, 2.6, 0.9])
    doc.add_page_break()

    # 2 dashboard
    feature(doc, "2", "Dashboard", FILL_NAVY, TEAL,
            "A one-look view of the whole order book: KPIs, weekly bottleneck utilisation (bulk analysis of "
            "every week in one pass), a machine x week capacity heatmap, the weeks that need attention, the "
            "connected dataset, and the feature roadmap.",
            [("Bulk processing", "every week is analysed in a single pass to build the overview and heatmap."),
             ("Heatmap", "work-centre x week grid coloured by utilisation - the classic capacity view."),
             ("Interactive", "click any red/amber week (chart or heatmap) to drill into it.")],
            "Foundation for 'capacity bottlenecks identified' and 'scheduling conflicts highlighted'.",
            "ss_1_dashboard.png", "Dashboard - KPIs, weekly overview, capacity heatmap, roadmap.")

    # 3 capacity
    feature(doc, "3", "Capacity Analysis & Bottleneck Detection", FILL_BLUE, TEAL,
            "For a selected week, the exact load on every machine (hours needed vs available), the bottleneck, "
            "an offload recommendation, the batch/campaign setup saving, and the orders due that week.",
            [("Method", "sum setup + run x qty per routing operation onto each machine; compare to weekly capacity."),
             ("Bottleneck", "highest-utilisation machine (Theory of Constraints)."),
             ("Batch insight", "setup hours saved if same-item orders run as one campaign.")],
            "Capacity bottlenecks identified correctly; scheduling conflicts (overloads) highlighted.",
            "ss_2_capacity.png", "Capacity Analysis - needed vs available, bottleneck, recommendation, batching, orders.")

    # 4 priority
    feature(doc, "4", "Order Prioritization", FILL_GREEN, GREEN,
            "Ranks the week's orders by urgency and explains why - so the planner knows exactly what to run first.",
            [("Drivers", "Earliest Due Date + Critical Ratio (time vs work) + customer tier + penalty."),
             ("At-risk flag", "critical ratio below 1 means not enough production time before the due date."),
             ("Transparent", "every row shows its score and the reasons behind its rank.")],
            "Supports order prioritization and feeds scheduling & scenarios.",
            "ss_3_priority.png", "Order Prioritization - ranked list with tier badges, critical-ratio warnings, urgency bars.")

    # 5 allocation
    feature(doc, "5", "Resource Allocation", FILL_BLUE, TEAL,
            "Moves work off overloaded machines onto qualified idle backups and shows the before/after balance.",
            [("Movable only", "only operations whose routing lists an alternate machine can move."),
             ("Safe", "never pushes a target machine over its own capacity."),
             ("Honest", "if a machine has no alternate, it says so (use overtime/outsourcing/deferral).")],
            "Alternative scheduling recommendations provided.",
            "ss_4_allocation.png", "Resource Allocation - offload move with before/after utilisation per machine.")

    # 6 schedule
    feature(doc, "6", "Schedule Optimization (Gantt)", FILL_MINT, GREEN,
            "Builds a feasible, near-optimal machine schedule with a constraint solver (Google OR-Tools CP-SAT) "
            "for the most urgent orders, and draws it as an interactive Gantt.",
            [("Model", "each order's operations run in routing order (precedence); one job per machine (no overlap)."),
             ("Objective", "minimise makespan - finish everything as early as possible."),
             ("Result", "solved to OPTIMAL in under a second; every block is a real operation.")],
            "Optimized production plan generated; scheduling conflicts resolved (no-overlap).",
            "ss_5_schedule.png", "Schedule Optimization - OPTIMAL Gantt, one row per machine, coloured per order.")

    # 7 delay risk
    feature(doc, "7", "Delay Risk", FILL_AMBER, AMBER,
            "Flags orders likely to be late - from missing materials (BOM vs inventory) or tight capacity - each "
            "with the root cause and a fix.",
            [("Material check", "explode each order's BOM, net weekly component demand against inventory (on-hand + on-order in time)."),
             ("Capacity check", "reuse the critical ratio; below 1 means at risk."),
             ("Actionable", "each at-risk order lists the cause and a concrete fix.")],
            "At-risk orders identified.",
            "ss_6_risk.png", "Delay Risk - at-risk orders with cause & fix, plus the short-components list.")

    # 8 demand
    feature(doc, "8", "Demand vs Capacity", FILL_TEAL, TEAL,
            "Answers the horizon-level question: can we commit to the whole order book? Total work required vs "
            "total capacity, by department and machine, with an honest verdict.",
            [("Aggregate", "sum required hours across all 16 weeks vs available; per department and machine."),
             ("Honest nuance", "committable in total does not mean every week fits - it flags the weekly peaks."),
             ("Fix", "cheapest way to close any gap, or which peak weeks to smooth.")],
            "Demand vs capacity assessment.",
            "ss_7_demand.png", "Demand vs Capacity - verdict, department chart, per-machine load, fixes.")

    # 9 scenarios
    feature(doc, "9", "What-if Scenarios", FILL_GREEN, GREEN,
            "Compares planning options for a week side by side, before committing to one.",
            [("Baseline", "the week as-is."),
             ("Add a shift", "give the bottleneck's department one extra shift."),
             ("Defer orders", "push the least-urgent orders until the bottleneck is within capacity."),
             ("Best tag", "marks the option that clears the most overload.")],
            "Supports minimum 3 planning scenarios.",
            "ss_8_scenarios.png", "What-if Scenarios - baseline vs add-a-shift vs defer, side by side.")

    # 10 copilot
    section(doc, "10", "AI Copilot (Azure OpenAI)", fill=FILL_NAVY)
    para(doc,
         "A floating assistant answers plain-language questions. It is powered by Azure OpenAI with real "
         "tool-calling: the model calls the deterministic planning tools (capacity, bottleneck, priority, "
         "allocation, delay-risk, demand, scenarios) and explains the result. Every number comes from the "
         "engine - the model never invents figures. If no key is configured, it falls back to a deterministic "
         "template so the app always works.")
    subhead(doc, "Example (live Azure OpenAI answer)")
    mono(doc, [
        "Q: Which week is most overloaded and what should I do about it?",
        "",
        "A: The most overloaded week is 2026-09-14, with 3 machines overloaded:",
        "   CNC-01 (114%), CAST-01 (111%), GRIND-01 (110%). CNC-01 is the bottleneck.",
        "   1) Reallocate 17h from CNC-01 to CNC-02 (brings CNC-01 to 100%).",
        "   2) Expedite two short components (Handwheel, Ball) for 7 at-risk orders.",
        "   3) Best option: defer 8 least-urgent orders -> 0 overloads.",
    ])
    shot(doc, "assistant-open2.png", width_in=3.6, caption="The copilot panel (bottom-right of every screen).")
    doc.add_page_break()

    # 11 architecture + tech + security
    section(doc, "11", "Architecture, Data & Security", fill=FILL_NAVY)
    arch = os.path.join(ROOT, "others", "architecture_diagram.png")
    if os.path.exists(arch):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(arch, width=Inches(6.4))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cap, "Deterministic engine at the core; AI copilot on top.", size=8.8, italic=True, color=SLATE)
    subhead(doc, "Technology")
    table(doc, ["Layer", "Technology"],
          [["Backend API", "FastAPI + Uvicorn (Python 3.12)"],
           ["Optimization", "Google OR-Tools CP-SAT"],
           ["AI copilot", "Azure OpenAI (chat completions + tool-calling)"],
           ["Frontend", "React + Vite + Recharts"],
           ["Data", "8 connected CSV tables (work_centers, items, routings, bom, components, inventory, customers, orders)"],
           ["Tests", "pytest (10 tests)"]],
          widths=[1.7, 4.6])
    subhead(doc, "Connected dataset (machines = work_centers.csv)")
    table(doc, ["Table", "Rows", "Note"],
          [["work_centers (machines)", "9", "CAST/CNC-01..03/GRIND/ASSY-01..02/TEST/PACK + capacity"],
           ["items", "45", "valve products"],
           ["routings", "270", "operations per item x machine"],
           ["bom", "330", "components per item"],
           ["components", "45", "parts + lead times"],
           ["inventory", "45", "stock per component"],
           ["customers", "30", "tier + penalty"],
           ["orders", "520", "the demand"]],
          widths=[2.2, 0.8, 3.5])
    subhead(doc, "Security review")
    bullets(doc, [
        ("Secrets", "the Azure key lives only in app/backend/.env; a .gitignore now excludes .env, .venv, logs, node_modules and dist. No git repo, so nothing has leaked."),
        ("Copilot safety", "all AI tools are READ-ONLY and take no arguments, so prompt injection cannot trigger any write or destructive action."),
        ("Input validation", "week parameters are validated (bad -> 400, unknown -> 404); schedule size is bounded."),
        ("Frontend", "0 production vulnerabilities (npm audit); data is rendered as text (no HTML injection)."),
        ("For production", "add authentication and lock CORS to the real origin (currently open to localhost for dev)."),
    ])
    doc.add_page_break()

    # 12 run
    section(doc, "12", "How to run", fill=FILL_TEAL)
    mono(doc, [
        "# Backend",
        "cd app/backend",
        "python -m venv .venv",
        r".\.venv\Scripts\python.exe -m pip install -r requirements.txt",
        r".\.venv\Scripts\python.exe -m uvicorn main:app --port 8001",
        "",
        "# Frontend",
        "cd app/frontend",
        "npm install",
        "npm run dev -- --port 5199   # open http://localhost:5199",
        "",
        "# Azure OpenAI copilot: fill app/backend/.env (AZURE_OPENAI_ENDPOINT / _API_KEY / _DEPLOYMENT)",
    ])
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    _run(p, "Feasible by construction. Explainable by design. Verified end to end.",
         size=11.5, italic=True, bold=True, color=NAVY)

    out = os.path.join(ROOT, "Production_Planning_Agent_FEATURES.docx")
    doc.save(out); print("Saved:", out)


if __name__ == "__main__":
    build()
